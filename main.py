import os
import shutil
from faker import Faker
from fpdf import FPDF
from docx import Document
import xlsxwriter
import pandas as pd
import random

# Create 'synthetic_files' directory if it doesn't exist
output_dir = 'synthetic_files'
if os.path.exists(output_dir):
    shutil.rmtree(output_dir)
os.makedirs(output_dir)

# Read lines from a file
def read_lines_from_file(file_path):
    with open(file_path, 'r') as file:
        lines = [line.strip() for line in file.readlines()]
    return lines

# Load machine names, product names, and sentence templates
machine_names_file = 'data/machine_names.txt'
product_names_file = 'product_names.txt'
sentence_templates_file = 'sentence_templates.txt'

realistic_machine_names = read_lines_from_file(machine_names_file)
realistic_product_names = read_lines_from_file(product_names_file)
sentence_templates = read_lines_from_file(sentence_templates_file)

# Initialize Faker and other setup
fake = Faker()
num_docs_per_type = 20  # Number of documents to create per type
num_unique_machines = len(realistic_machine_names)
num_unique_products = len(realistic_product_names)
test_results = []

# Function to generate fake test results with random machine names and products
def generate_test_result():
    machine_name = random.choice(realistic_machine_names)
    product = random.choice(realistic_product_names)
    result = round(random.uniform(0, 100), 2)
    return machine_name, product, result

# Generate synthetic test results for each document
doc_test_results_list = []
for _ in range(num_docs_per_type * 3):  # 3 types: Word, Excel, PDF
    num_results = random.randint(1, 20)  # Random number of test results per document
    doc_test_results = [generate_test_result() for _ in range(num_results)]
    doc_test_results_list.append(doc_test_results)
    test_results.extend(doc_test_results)

# Save the master file with all test results
master_file_path = os.path.join(output_dir, 'master_test_results.csv')
master_df = pd.DataFrame(test_results, columns=['Machine Name', 'Product', 'Result'])
master_df.to_csv(master_file_path, index=False)

# Function to create a realistic report-like structure with random paragraph and table placement
def create_report_content(test_results):
    content = []
    remaining_results = test_results.copy()

    while remaining_results:
        # Randomly decide the number of paragraphs and results to add
        num_paragraphs = random.randint(1, 3)
        num_results = min(len(remaining_results), random.randint(1, 5))
        add_table = random.choice([True, False])
        add_noise = random.choice([True, False])

        # Add random paragraphs
        for _ in range(num_paragraphs):
            paragraph = fake.paragraph(nb_sentences=random.randint(2, 5))
            content.append(paragraph)
            # Add bullet points or numbered lists occasionally
            if random.choice([True, False]):
                list_type = random.choice(['bullet', 'number'])
                list_items = [fake.sentence() for _ in range(random.randint(3, 5))]
                if list_type == 'bullet':
                    content.append(('bullet', list_items))
                else:
                    content.append(('number', list_items))

        # Add noise text
        if add_noise:
            noise_paragraph = fake.paragraph(nb_sentences=random.randint(1, 3))
            content.append(noise_paragraph)

        # Add test results either as text or in a table
        if add_table and len(remaining_results) >= num_results:
            table_results = remaining_results[:num_results]
            remaining_results = remaining_results[num_results:]
            content.append(('table', table_results))
        else:
            for _ in range(num_results):
                if remaining_results:
                    machine_name, product, result = remaining_results.pop(0)
                    template = random.choice(sentence_templates)
                    content.append(template.format(machine=machine_name, product=product, result=result))

    return content

# Function to create multiple Word documents
def create_multiple_word_docs(num_docs, doc_test_results_list):
    for i in range(num_docs):
        file_name = os.path.join(output_dir, f'synthetic_test_results_{i+1}.docx')
        doc = Document()
        doc.add_heading('Synthetic Test Report', level=1)
        doc_test_results = doc_test_results_list[i]
        report_content = create_report_content(doc_test_results)
        for item in report_content:
            if isinstance(item, tuple):
                if item[0] == 'table':
                    table = doc.add_table(rows=1, cols=3)
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Machine Name'
                    hdr_cells[1].text = 'Product'
                    hdr_cells[2].text = 'Test Result'
                    for machine_name, product, result in item[1]:
                        row_cells = table.add_row().cells
                        row_cells[0].text = machine_name
                        row_cells[1].text = product
                        row_cells[2].text = str(result)
                elif item[0] == 'bullet':
                    for list_item in item[1]:
                        p = doc.add_paragraph(list_item)
                        p.style = 'List Bullet'
                elif item[0] == 'number':
                    for list_item in item[1]:
                        p = doc.add_paragraph(list_item)
                        p.style = 'List Number'
            else:
                doc.add_paragraph(item)
        doc.save(file_name)

# Function to create multiple Excel files
def create_multiple_excel_files(num_docs, doc_test_results_list):
    for i in range(num_docs):
        file_name = os.path.join(output_dir, f'synthetic_test_results_{i+1}.xlsx')
        workbook = xlsxwriter.Workbook(file_name)
        worksheet = workbook.add_worksheet()
        row = 0
        doc_test_results = doc_test_results_list[num_docs + i]
        report_content = create_report_content(doc_test_results)
        for item in report_content:
            if isinstance(item, tuple) and item[0] == 'table':
                worksheet.write(row, 0, 'Machine Name')
                worksheet.write(row, 1, 'Product')
                worksheet.write(row, 2, 'Result')
                row += 1
                for machine_name, product, result in item[1]:
                    worksheet.write(row, 0, machine_name)
                    worksheet.write(row, 1, product)
                    worksheet.write(row, 2, result)
                    row += 1
            elif isinstance(item, tuple):
                # For lists, write them in a merged cell
                if item[0] == 'bullet' or item[0] == 'number':
                    worksheet.write(row, 0, '\n'.join(item[1]))
                    row += 1
            else:
                worksheet.write(row, 0, item)
                row += 1
        workbook.close()

# Function to create multiple PDF files
def create_multiple_pdf_files(num_docs, doc_test_results_list):
    for i in range(num_docs):
        file_name = os.path.join(output_dir, f'synthetic_test_results_{i+1}.pdf')
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt="Synthetic Test Report", ln=True, align='C')
        doc_test_results = doc_test_results_list[2 * num_docs + i]
        report_content = create_report_content(doc_test_results)
        for item in report_content:
            if isinstance(item, tuple) and item[0] == 'table':
                pdf.set_font("Arial", size=10)
                pdf.cell(40, 10, 'Machine Name', 1)
                pdf.cell(40, 10, 'Product', 1)
                pdf.cell(40, 10, 'Test Result', 1)
                pdf.ln()
                for machine_name, product, result in item[1]:
                    pdf.cell(40, 10, machine_name, 1)
                    pdf.cell(40, 10, product, 1)
                    pdf.cell(40, 10, str(result), 1)
                    pdf.ln()
                pdf.set_font("Arial", size=12)
            elif isinstance(item, tuple):
                if item[0] == 'bullet' or item[0] == 'number':
                    list_content = '\n'.join(item[1])
                    pdf.multi_cell(0, 10, txt=list_content)
            else:
                pdf.multi_cell(0, 10, txt=item)
        pdf.output(file_name)

# Create multiple documents
create_multiple_word_docs(num_docs_per_type, doc_test_results_list)
create_multiple_excel_files(num_docs_per_type, doc_test_results_list)
create_multiple_pdf_files(num_docs_per_type, doc_test_results_list)

print("Multiple synthetic files and master CSV file have been created in the 'synthetic_files' directory.")
