import os
import random
import time
from datetime import datetime
from io import StringIO
import markdown
import numpy as np
from bs4 import BeautifulSoup
from openai import AzureOpenAI
from dotenv import load_dotenv
import re
import csv
from fpdf import FPDF
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from xlsxwriter.utility import xl_range


# Load environment variables
load_dotenv()

# Azure OpenAI model and client configuration
model_name = "gpt-4o-2024-08-06"
client = AzureOpenAI(
    azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
    api_key=os.getenv("AZURE_OPENAI_API_KEY"),
    api_version="2024-02-01"
)

# Set number of files to be created
num_files = 2500
num_model_pts = 5000

# ------------------------------------------------------------------------------------

# Define the machine list with only the Viscometer
MACHINE = "Viscometer VS-300"

# Define ingredients and clusters
BASE_OILS = ["Jojoba Oil", "Coconut Oil", "Almond Oil"]
THICKENERS = ["Beeswax", "Gum", "Cetyl Alcohol"]
ADDITIVES = ["Vitamin E", "Glycerin"]

# Define the product list with combinations of base oils, thickeners, and additives
full_product_list = [
    f"{oil}" for oil in BASE_OILS
] + [
    f"{oil}, {thickener}" for oil in BASE_OILS for thickener in THICKENERS
] + [
    f"{oil}, {additive}" for oil in BASE_OILS for additive in ADDITIVES
] + [
    f"{oil}, {thickener}, {additive}" for oil in BASE_OILS for thickener in THICKENERS for additive in ADDITIVES
]

# Randomly select 6 combinations to exclude
excluded_combinations = random.sample(full_product_list, 6)

# Create the final product list, excluding the selected combinations
PRODUCT_LIST = [product for product in full_product_list if product not in excluded_combinations]

CLUSTERS = {
    "base_oil": {
        "Jojoba Oil": 2500,
        "Coconut Oil": 5000,
        "Almond Oil": 7500
    },
    "thickener": {
        "Beeswax": {
            "Jojoba Oil": 500,
            "Coconut Oil": -150,
            "Almond Oil": -350
        },
        "Gum": {
            "Jojoba Oil": -500,
            "Coconut Oil": 250,
            "Almond Oil": 125
        },
        "Cetyl Alcohol": {
            "Jojoba Oil": 300,
            "Coconut Oil": 100,
            "Almond Oil": -250
        }
    },
    "additive": {
        "Vitamin E": {
            "Jojoba Oil": 50,
            "Coconut Oil": -75,
            "Almond Oil": 25
        },
        "Glycerin": {
            "Jojoba Oil": -100,
            "Coconut Oil": 25,
            "Almond Oil": 50
        }
    }
}

def assign_viscometer_result(product):
    """
    Assigns a synthetic result for the Viscometer based on the specific effects of each ingredient.
    """
    result = 0
    base_oil = None

    # Determine the base oil in the product
    for oil in BASE_OILS:
        if oil in product:
            base_oil = oil
            result += CLUSTERS['base_oil'][base_oil]
            break

    # Adjust the result based on thickeners and additives in relation to the base oil
    if base_oil:
        for ingredient in product.split(', '):
            if ingredient in CLUSTERS['thickener'] and base_oil in CLUSTERS['thickener'][ingredient]:
                result += CLUSTERS['thickener'][ingredient][base_oil]
            elif ingredient in CLUSTERS['additive'] and base_oil in CLUSTERS['additive'][ingredient]:
                result += CLUSTERS['additive'][ingredient][base_oil]

    # Add some noise to simulate realistic data
    result += np.random.normal(0, 100)
    return max(1, round(result, 2))  # Ensure results stay within a realistic range

def generate_data_points(num_samples=1):
    """
    Generates unique data points for the Viscometer formatted for clustering and predictive modeling.
    """
    data = []

    for _ in range(num_samples):
        # Machine is always Viscometer VS-300
        machine = MACHINE

        # Randomly select a product mixture
        product = np.random.choice(PRODUCT_LIST)

        # Split product into individual ingredients
        ingredients = product.split(', ')

        # Ensure exactly three ingredients are represented in the columns
        ing_1 = ingredients[0] if len(ingredients) > 0 else ''
        ing_2 = ingredients[1] if len(ingredients) > 1 else ''
        ing_3 = ingredients[2] if len(ingredients) > 2 else ''

        # Generate the result based solely on the product composition
        result = assign_viscometer_result(product)
        data.append([machine, ing_1, ing_2, ing_3, result, 'cP'])

    # Create DataFrame with the new columns
    columns = ['machine', 'ing_1', 'ing_2', 'ing_3', 'result', 'unit']
    dataset = pd.DataFrame(data, columns=columns)

    return dataset

def save_dataset(dataset, folder_path):
    """
    Saves the viscosity dataset as a CSV file in the specified folder.

    Parameters:
    - dataset (DataFrame): The DataFrame containing the viscosity data.
    - folder_path (str): The path to the folder where the dataset should be saved.
    """
    # Define the file path for the dataset CSV
    dataset_file_path = os.path.join(folder_path, 'a_viscosity_dataset.csv')

    # Save the dataset as a CSV file
    dataset.to_csv(dataset_file_path, index=False)
    print(f"Viscosity dataset saved as {dataset_file_path}")

# Generate the dataset
dataset = generate_data_points(num_samples=num_model_pts)
unused_data_indices = list(range(len(dataset)))  # Track unused indices

# ------------------------------------------------------------------------------------

def generate_document_name(report_counter):
    """
    Generates a document name based on the current report number.
    """
    return f"Report_{report_counter}"


def generate_openai_response(prompt):
    """
    Function to generate a response from the Azure OpenAI API.

    Parameters:
    - prompt (str): The input text prompt for the model.

    Returns:
    - response_text (str): The generated text from the Azure OpenAI API.
    """
    try:
        # Call the Azure OpenAI API for chat completions
        response = client.chat.completions.create(
            model=model_name,
            messages=[
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": prompt}
            ]
        )

        # Extract and return the generated text
        response_text = response.choices[0].message.content
        return response_text

    except Exception as e:
        print(f"Error generating text: {e}")
        return None


def run_interaction_loop(file_num):
    # Initialize the personas
    orchestrator = OrchestratorPersona(name="Orchestrator")
    document_creator = DocumentCreatorPersona(name="Document Creator")

    # Counter to track the number of successfully generated files
    files_created = 0
    report_counter = 1

    # Run the loop to create documents, ensuring all data points are used
    while files_created < file_num:
        # Calculate remaining documents and ensure that all remaining data points will be used
        remaining_reports = file_num - files_created
        remaining_data_points = len(unused_data_indices)

        # Determine maximum data points to inject per report to ensure all are used
        max_inject_per_report = (remaining_data_points + remaining_reports - 1) // remaining_reports

        # Orchestrator creates a prompt for the creator with selected products and machines
        selected_products = random.sample(orchestrator.products, random.randint(5, 10))
        selected_machines = random.sample(orchestrator.machines, random.randint(5, 10))
        print('[Orchestrator] Creating Initial Prompts')

        # Generate the report name
        report_name = generate_document_name(report_counter)

        # Document Creator generates a report and an answer key based on the prompt
        report_content, answer_key = document_creator.generate_report_with_answer_key(
            selected_products, selected_machines, report_name, max_inject_per_report
        )
        print('[Creator] Generating Answer Key & Report')

        # Save the report directly without evaluation
        orchestrator.save_report(report_content, answer_key, report_name)
        files_created += 1
        report_counter += 1  # Increment the report counter
        print(f"Report {files_created} saved successfully.\n")

    # Inject any remaining data points if necessary
    while unused_data_indices:
        selected_products = random.sample(orchestrator.products, random.randint(5, 10))
        selected_machines = random.sample(orchestrator.machines, random.randint(5, 10))
        report_name = generate_document_name(report_counter)

        # Ensure that all remaining data points are used in subsequent reports
        report_content, answer_key = document_creator.generate_report_with_answer_key(
            selected_products, selected_machines, report_name, max_inject_per_report=5
        )
        orchestrator.save_report(report_content, answer_key, report_name)
        files_created += 1
        report_counter += 1
        print(f"Injecting remaining data points: Report {files_created} saved successfully.\n")

    # Save the viscosity dataset in the folder where reports are stored
    save_dataset(dataset, orchestrator.run_folder)


class DocumentCreatorPersona:
    def __init__(self, name):
        self.name = name

    def generate_report_with_answer_key(self, selected_products, selected_machines, report_name, max_inject_per_report):
        """
        Generates both the answer key and the report content using the Azure OpenAI API.
        """
        # Step 1: Generate the answer key using the LLM
        answer_key = self.generate_answer_key(selected_products, selected_machines, report_name)

        # Step 2: Randomly decide how many data points to inject (0 to max_inject_per_report)
        num_to_inject = min(max_inject_per_report, random.randint(0, 10))

        # Inject the selected number of data points
        for _ in range(num_to_inject):
            if unused_data_indices:  # Check if there are any data points left to inject
                selected_index = unused_data_indices.pop(0)  # Use the next available data point
                data_row = dataset.iloc[selected_index]

                # Extract individual fields from the data row
                machine = data_row['machine']
                ing_1 = data_row['ing_1']
                ing_2 = data_row['ing_2']
                ing_3 = data_row['ing_3']
                result = data_row['result']
                unit = data_row['unit']

                # Create the CSV-formatted line for the answer key with individual ingredients
                data_csv = f'{report_name},{machine},"{ing_1}","{ing_2}","{ing_3}",{result},{unit}'
                answer_key += f'\n{data_csv}'

        # Step 3: Generate the report content based on the updated answer key
        report_content = self.generate_report_content(answer_key)

        return report_content, answer_key

    def generate_answer_key(self, selected_products, selected_machines, report_name):
        """
        Generates an answer key using the Azure OpenAI API.
        """
        prompt = (
            "Generate an answer key for a lab report in CSV format with the structure: report name, machine, ing_1, ing_2, ing_3, result, unit. "
            "Ensure each row has at least one oil in ing_1, followed by optional ingredients in ing_2 and ing_3. "
            "Do not add headers, quotes, or any extra text. Use the exact machine names and units provided. "
            "Machines: " + ', '.join(selected_machines) + ". "
            "Ingredient combinations: " + ', '.join([f'"{p}"' for p in selected_products]) + ". "
            f"Fill the report name as {report_name} for all rows. "
            "Output only the requested CSV data without explanations or formatting."
        )

        # generate answer key
        answer_key = generate_openai_response(prompt)

        return answer_key

    def generate_report_content(self, answer_key):
        """
        Generates the report content using the Azure OpenAI API based on the provided answer key.
        """
        prompt = (
            f"Create a detailed lab report using the following answer key data:\n\n{answer_key}\n\n"
            "The report should treat each set of ingredients as a single test sample and generate results accordingly. "
            "Each group of ingredients (e.g., 'Jojoba Oil, Beeswax, Vitamin E') should be considered as a mixture and tested together. "
            "The report should include observations, measurements, results, and descriptions based on the provided answer key. "
            "Include multiple paragraphs, multiple tables with mixed data types, randomly scattered irrelevant information, "
            "and complex descriptions. Ensure the data is presented in a way that is challenging to extract via an automated process."
            "The information in the answer key must be somewhere in the report."
        )

        # Generate the report content
        report_content = generate_openai_response(prompt)
        return report_content


class DocumentEvaluatorPersona:
    def __init__(self, name):
        self.name = name

    def evaluate_document(self, document):
        """
        Evaluates the complexity of extracting data from a given document using the Azure OpenAI API.
        """
        prompt = (f"Evaluate the following lab report for its difficulty in extracting structured data. Consider "
                  f"factors such as mixed data formats, irrelevant information, inconsistent data, scattered facts, "
                  f"and complexity of language. Provide an extraction rating from 1 (easy) to 10 (hard)."
                  f"The rating should be the last line and in the format: 'Rating: X' where X is a number between 1 and 10.\n\n"
                  f"{document}")

        # Use the helper function to generate a response
        evaluation = generate_openai_response(prompt)

        # Check if the evaluation contains a valid rating
        if not re.search(r"Rating:\s*\d+", evaluation):
            print("[Evaluator] LLM response does not contain a valid rating. Retrying")
            evaluation = generate_openai_response(prompt)  # Retry the request

        return evaluation


class OrchestratorPersona:
    def __init__(self, name):
        self.name = name
        self.difficulty_level = 7
        self.products = self.load_data('data/product_names.txt')
        self.machines = self.load_data('data/machine_names.txt')

        # Create a single timestamped directory for all reports and answer key
        self.base_folder = 'synthetic_files'
        timestamp = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
        self.run_folder = os.path.join(self.base_folder, f"reports_{timestamp}")
        os.makedirs(self.run_folder, exist_ok=True)  # Create the directory

        # Initialize the answer key CSV file
        self.answer_key_file = os.path.join(self.run_folder, 'answer_key.csv')
        self.initialize_answer_key()

    def load_data(self, filename):
        """
        Loads data from a text file and returns a list of lines.
        """
        with open(filename, 'r') as file:
            return [line.strip() for line in file.readlines()]

    def evaluate_evaluator_report(self, evaluation, report_content, answer_key, report_name):
        """
        Analyzes the evaluator's feedback and decides whether to save or discard the report.
        """
        # Extract difficulty rating from the evaluation response using a flexible regular expression
        match = re.search(r"Rating:\s*(\d+)", evaluation, re.IGNORECASE)

        if match:
            difficulty_rating = int(match.group(1))  # Extract the first number found
            print('[Evaluator] Extraction Difficulty: ', difficulty_rating)

            if difficulty_rating >= 7:
                print(f"[{self.name}] Report is difficult enough. Saving report.")
                self.save_report(report_content, answer_key, report_name)
                return True
            else:
                print(f"[{self.name}] Report is not difficult enough. Increasing difficulty.")
                # Increase difficulty for the next prompt
                self.difficulty_level = difficulty_rating + 1
                return False
        else:
            print(f"[{self.name}] Could not find a valid rating in the evaluation. Evaluation output was: {evaluation}")
            return False

    def save_report(self, content, answer_key, report_name):
        """
        Saves the report in a chosen format (DOCX, XLSX, TXT, PDF) and appends answer key data
        to a single answer key CSV in a timestamped directory.
        """
        # Define the formats and their respective weights
        formats = ['xlsx', 'docx', 'txt', 'pdf']
        weights = [0.5, 0.3, 0.1, 0.1]  # Excel 50%, Word 30%, Text 10%, PDF 10%

        # Choose a format based on the defined probabilities
        chosen_format = random.choices(formats, weights, k=1)[0]

        # Update file path to include the new timestamped directory
        file_name = os.path.join(self.run_folder, f"{report_name}.{chosen_format}")

        # Save the report in the chosen format
        paragraphs, tables = self.parse_content(content)

        try:
            if chosen_format == 'docx':
                try:
                    document = content

                    # Convert Markdown to HTML with tables support
                    html = markdown.markdown(document, extensions=['tables'])

                    # Parse the HTML with BeautifulSoup
                    soup = BeautifulSoup(html, 'html.parser')

                    # Extract paragraphs and tables in the order they appear
                    content_list = []

                    # Combine all paragraphs and tables together by iterating through the soup's children
                    for elem in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'table']):
                        if elem.name == 'table':
                            # Process tables
                            table_str = StringIO(str(elem))
                            df = pd.read_html(table_str)[0]
                            content_list.append((df, 'table'))
                        else:
                            # Process paragraphs and headings
                            text = elem.get_text(strip=True)
                            if text:
                                content_list.append((text, 'paragraph'))

                    # Initialize the Word document
                    doc = Document()

                    # Function to add paragraphs to the Word document
                    def add_paragraph_to_doc(doc, paragraph):
                        doc.add_paragraph(paragraph)

                    # Function to add tables to the Word document
                    def add_table_to_doc(doc, df):
                        table_doc = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
                        table_doc.style = 'Table Grid'

                        # Add header row
                        hdr_cells = table_doc.rows[0].cells
                        for j, column_name in enumerate(df.columns):
                            hdr_cells[j].text = str(column_name)
                            hdr_cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            hdr_cells[j].paragraphs[0].runs[0].bold = True

                        # Add data rows
                        for i, row in df.iterrows():
                            row_cells = table_doc.rows[i + 1].cells
                            for j, cell in enumerate(row):
                                row_cells[j].text = str(cell)
                                row_cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Add content to the Word document in the original order
                    for item, item_type in content_list:
                        if item_type == 'paragraph':
                            add_paragraph_to_doc(doc, item)
                        elif item_type == 'table':
                            add_table_to_doc(doc, item)

                    # Save the Word document
                    doc.save(file_name)
                    print(f"Report saved as a Word document: {file_name}")

                except Exception as e:
                    print(f"Error saving DOCX file: {e}")

            elif chosen_format == 'xlsx':
                try:
                    document = content

                    # Convert Markdown to HTML
                    html = markdown.markdown(document, extensions=['tables'])

                    # Parse the HTML with BeautifulSoup
                    soup = BeautifulSoup(html, 'html.parser')

                    # Remove bold tags from the HTML (optional)
                    for bold_tag in soup.find_all(['strong', 'b']):
                        bold_tag.replace_with(bold_tag.text)

                    # Extract paragraphs
                    paragraphs = []
                    for elem in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                        text = elem.get_text(strip=True)
                        if text:
                            paragraphs.append(text)

                    # Combine all paragraphs into a single text block
                    combined_paragraphs = "\n".join(paragraphs)

                    # Extract tables
                    tables = []
                    html_tables = soup.find_all('table')
                    for table in html_tables:
                        table_str = StringIO(str(table))
                        df = pd.read_html(table_str)[0]
                        tables.append(df)

                    with pd.ExcelWriter(file_name, engine='xlsxwriter') as writer:
                        workbook = writer.book

                        # Save combined paragraphs into a merged cell
                        if combined_paragraphs:
                            worksheet = workbook.add_worksheet('Report')

                            # Define the range to merge: 10 columns wide and 20 rows long
                            merge_range = xl_range(0, 0, 29, 19)
                            worksheet.merge_range(merge_range, combined_paragraphs)

                        else:
                            print("No paragraphs to save.")

                        # Save tables
                        if tables:
                            for idx, df_table in enumerate(tables):
                                sheet_name = f'Table_{idx + 1}'
                                df_table.to_excel(writer, sheet_name=sheet_name, index=False)
                        else:
                            print("No tables to save.")

                    print(f"Report saved as an Excel file: {file_name}")

                except Exception as e:
                    print(f"Error saving XLSX file: {e}")

            elif chosen_format == 'txt':
                try:
                    with open(file_name, 'w', encoding='utf-8') as txt_file:
                        for para in paragraphs:
                            txt_file.write(para + '\n\n')
                        txt_file.write("\nTables:\n")
                        for table in tables:
                            for row in table:
                                txt_file.write('\t'.join(row) + '\n')
                    print(f"Report saved as a TXT file: {file_name}")

                except Exception as e:
                    print(f"Error saving TXT file: {e}")

            elif chosen_format == 'pdf':
                try:
                    document = content

                    # Convert Markdown to HTML with tables support
                    html = markdown.markdown(document, extensions=['tables'])

                    # Parse the HTML with BeautifulSoup
                    soup = BeautifulSoup(html, 'html.parser')

                    # Extract paragraphs and tables in the order they appear
                    content_list = []

                    # Combine all paragraphs and tables together by iterating through the soup's children
                    for elem in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'table']):
                        if elem.name == 'table':
                            # Process tables
                            table_str = StringIO(str(elem))
                            df = pd.read_html(table_str)[0]
                            content_list.append((df, 'table'))
                        else:
                            # Process paragraphs and headings
                            text = elem.get_text(strip=True)
                            if text:
                                content_list.append((text, 'paragraph'))

                    # Initialize PDF
                    pdf = FPDF()
                    pdf.set_auto_page_break(auto=True, margin=15)
                    pdf.add_page()
                    pdf.set_font("Arial", size=10)

                    # Function to add paragraphs to the PDF
                    def add_paragraph_to_pdf(pdf, paragraph):
                        pdf.multi_cell(0, 10, paragraph.encode('latin-1', 'replace').decode('latin-1'))
                        pdf.ln(5)  # Add a new line after each paragraph

                    # Function to add tables to the PDF
                    def add_table_to_pdf(pdf, df):
                        page_width = pdf.w - 2 * pdf.l_margin
                        col_width = page_width / len(df.columns)  # Set equal width for all columns

                        # Add header row
                        pdf.set_font("Arial", 'B', 10)
                        for col_name in df.columns:
                            pdf.cell(col_width, 10, col_name, border=1, align='C')
                        pdf.ln(10)

                        # Add data rows
                        pdf.set_font("Arial", size=10)
                        for index, row in df.iterrows():
                            for cell in row:
                                cell_text = str(cell)
                                pdf.cell(col_width, 10, cell_text, border=1, align='C')
                            pdf.ln(10)

                    # Add content to the PDF in the original order
                    for item, item_type in content_list:
                        if item_type == 'paragraph':
                            add_paragraph_to_pdf(pdf, item)
                        elif item_type == 'table':
                            add_table_to_pdf(pdf, item)

                    # Save the PDF document
                    pdf.output(file_name)
                    print(f"Report saved as a PDF file: {file_name}")

                except Exception as e:
                    print(f"Error saving PDF file: {e}")

        except Exception as e:
            print(f"Error during report saving process: {e}")

            # Append answer key to the single CSV file
        try:
            self.append_answer_key_to_csv(answer_key)
        except Exception as e:
            print(f"Error appending to CSV: {e}")

    def initialize_answer_key(self):
        """
        Initializes the answer key CSV file with headers if it does not already exist.
        """
        if not os.path.exists(self.answer_key_file):
            with open(self.answer_key_file, mode='w', newline='') as file:
                writer = csv.writer(file)
                writer.writerow(['report name', 'machine', 'ing_1', 'ing_2', 'ing_3', 'result', 'unit'])

    def append_answer_key_to_csv(self, answer_key):
        """
        Appends the generated answer key to the single 'answer_key.csv' file.
        Cleans up unnecessary text artifacts such as descriptions, headers, or unwanted characters.
        """
        # Convert the answer key from text to structured data (assuming LLM outputs CSV format)
        answer_key_lines = answer_key.strip().split('\n')

        # List to hold the cleaned lines
        cleaned_lines = []
        header_written = False  # Flag to check if header has been written already

        # Check if the header exists in the current CSV file
        if os.path.exists(self.answer_key_file):
            with open(self.answer_key_file, 'r') as file:
                existing_lines = file.readlines()
                header_written = any(
                    'report name,machine,ing_1,ing_2,ing_3,result,unit' in line.lower() for line in existing_lines)

        # Iterate through the lines to clean up the content
        for line in answer_key_lines:
            # Skip lines that contain non-CSV content or explanations
            if any(keyword in line.lower() for keyword in ["format", "includes", "columns", "output"]):
                continue

            # Skip any lines that are blank
            if line.strip() == "":
                continue

            # Check for headers and ensure they are not duplicated
            if line.lower().startswith('report name,machine,ing_1,ing_2,ing_3,result,unit'):
                if header_written:
                    continue
                else:
                    header_written = True

            # Split line parts using regex to handle quoted sections and commas correctly
            line_parts = [part.strip().strip('"') for part in
                          re.split(r',(?=(?:[^\"]*\"[^\"]*\")*[^\"]*$)', line)]  # Strip quotes from each part

            # Skip any lines that don't have enough data
            if len(line_parts) < 4:
                continue

            # Extract result and unit from the end
            result = line_parts[-2] if len(line_parts) >= 2 else ''
            unit = line_parts[-1] if len(line_parts) >= 1 else ''

            # Extract report name and machine
            report_name = line_parts[0]
            machine = line_parts[1]

            # Extract ingredients from positions after machine and before result and unit
            ingredients = line_parts[2:-2]

            # Pad the ingredients list to ensure it has exactly 3 elements
            while len(ingredients) < 3:
                ingredients.append('')

            # Reconstruct line_parts
            line_parts = [report_name, machine] + ingredients[:3] + [result, unit]

            cleaned_lines.append(line_parts)

        # Append to CSV
        with open(self.answer_key_file, mode='a', newline='') as file:
            writer = csv.writer(file)
            # Write the header only once if it hasn't been written
            if not header_written:
                writer.writerow(['report name', 'machine', 'ing_1', 'ing_2', 'ing_3', 'result', 'unit'])
                header_written = True
            for line in cleaned_lines:
                # Extra check to ensure no unwanted characters remain
                if "'''" not in line and "```" not in line:
                    writer.writerow(line)

    def parse_content(self, content):
        """
        Parses the LLM-generated content to separate paragraphs, tables, and chart descriptions.
        """
        paragraphs = []
        tables = []

        # Simple parsing logic (this can be expanded based on LLM output format)
        lines = content.split('\n')
        current_table = []
        for line in lines:
            if 'Table:' in line:
                if current_table:
                    tables.append(current_table)
                    current_table = []
                current_table.append(line.split('\t'))
            else:
                paragraphs.append(line)

        if current_table:
            tables.append(current_table)

        return paragraphs, tables


if __name__ == "__main__":
    # Start the timer before calling the function
    start_time = time.time()

    # Run the interaction loop
    run_interaction_loop(num_files)

    # Stop the timer after the function completes
    end_time = time.time()

    # Calculate and print the total runtime
    total_runtime = end_time - start_time
    print(f"Total time to generate all files: {total_runtime:.2f} seconds.")
    print(f"Average time per file: {total_runtime / num_files:.2f} seconds.")
